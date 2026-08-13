import io
import json
import re
import csv
import zipfile
import hmac
import hashlib
from pathlib import Path

import numpy as np
import pandas as pd
import streamlit as st
import streamlit.components.v1 as components


# ============================================================
# إعداد الصفحة
# ============================================================
st.set_page_config(
    page_title="منصة تحليلات الجودة التعليمية",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ضبط لغة واتجاه المستند لقارئات الشاشة؛ CSS وحده لا يغيّر خصائص HTML الجذرية.
components.html(
    """
    <script>
      const root = window.parent.document.documentElement;
      root.setAttribute("lang", "ar");
      root.setAttribute("dir", "rtl");
    </script>
    """,
    height=0,
)

st.markdown(
    """
    <style>
    :root {
        --primary: #79b99a;
        --primary-strong: #4f9274;
        --primary-dark: #235c48;
        --primary-soft: #e8f5ed;
        --secondary: #c9a24b;
        --secondary-dark: #8a691f;
        --secondary-soft: #fff7e3;
        --ink: #173d31;
        --muted: #5d746b;
        --surface: #ffffff;
        --soft: #f4faf6;
        --line: #d8e8de;
        --shadow: 0 12px 32px rgba(35,92,72,.08);
    }

    html, body, [class*="css"], [data-testid="stAppViewContainer"] {
        font-family: "Tajawal", "Segoe UI", Tahoma, sans-serif;
    }
    [data-testid="stAppViewContainer"] {
        direction: rtl;
        background:
          radial-gradient(circle at 8% 4%, rgba(201,162,75,.14), transparent 25rem),
          radial-gradient(circle at 92% 18%, rgba(121,185,154,.16), transparent 28rem),
          linear-gradient(180deg, #f5fbf7 0%, #ffffff 42%, #fbfdfb 100%);
        color: var(--ink);
    }
    [data-testid="stMainBlockContainer"] {max-width: 1320px; padding-top: 1.6rem; padding-bottom:2rem;}
    [data-testid="stSidebar"] {direction: rtl; border-left: 1px solid #cfdfd5;}
    [data-testid="stSidebar"] > div {
        background:linear-gradient(180deg,#edf7f1 0%,#fffaf0 100%);
    }
    h1, h2, h3 {color: var(--ink); letter-spacing: -.02em;}
    h2 {
        position:relative; border-right:5px solid var(--secondary); padding:.15rem .8rem .2rem 0;
        margin-top:2.15rem; font-weight:850;
    }
    h2:after {
        content:""; display:block; width:54px; height:3px; border-radius:99px;
        background:var(--primary); margin-top:.5rem;
    }

    .hero {
        position:relative; isolation:isolate; overflow:hidden; color:var(--ink);
        background:linear-gradient(125deg,#f2fbf6 0%,#cfeada 46%,#8dc5a8 100%);
        border:1px solid rgba(79,146,116,.25); border-radius:28px;
        padding:2.55rem 2.65rem; margin-bottom:1.45rem;
        box-shadow:0 20px 50px rgba(35,92,72,.14);
    }
    .hero:before {
        content:""; position:absolute; inset:0 auto 0 0; width:10px;
        background:linear-gradient(180deg,var(--secondary),#e4c779); z-index:-1;
    }
    .hero:after {
        content:""; position:absolute; width:300px; height:300px; left:-85px; top:-125px;
        border:48px solid rgba(255,255,255,.3); border-radius:50%; z-index:-1;
    }
    .hero-kicker {
        display:inline-flex; align-items:center; gap:.4rem; font-size:.78rem; font-weight:850;
        letter-spacing:.08em; color:var(--secondary-dark); background:rgba(255,247,227,.85);
        border:1px solid rgba(201,162,75,.45); padding:.35rem .7rem; border-radius:99px;
        margin-bottom:.8rem;
    }
    .hero h1 {color:#174c3a; margin:0 0 .65rem; font-size:clamp(2rem,4vw,3.2rem); line-height:1.16; font-weight:900;}
    .hero p {margin:0; color:#315f50; font-size:1.08rem; max-width:790px; line-height:1.85;}
    .hero-badges {display:flex; gap:.55rem; flex-wrap:wrap; margin-top:1.2rem;}
    .hero-badge {
        background:rgba(255,255,255,.72); color:#245b48; border:1px solid rgba(79,146,116,.27);
        padding:.42rem .78rem; border-radius:99px; font-size:.82rem; font-weight:750;
        box-shadow:0 4px 12px rgba(35,92,72,.05);
    }

    .feature-grid {display:grid; grid-template-columns:repeat(3,1fr); gap:1rem; margin:.9rem 0 1.4rem;}
    .feature-card {
        position:relative; overflow:hidden; background:rgba(255,255,255,.92);
        border:1px solid var(--line); border-radius:19px; padding:1.15rem 1.2rem;
        box-shadow:var(--shadow); transition:transform .2s ease,box-shadow .2s ease;
    }
    .feature-card:before {content:""; position:absolute; top:0; right:0; width:100%; height:4px; background:linear-gradient(90deg,var(--primary),var(--secondary));}
    .feature-card:hover {transform:translateY(-3px); box-shadow:0 16px 34px rgba(35,92,72,.12);}
    .feature-card b {display:block; color:var(--primary-dark); margin-bottom:.3rem; font-size:.98rem;}
    .feature-card span {font-size:.87rem; color:var(--muted); line-height:1.7;}

    [data-testid="stFileUploader"] {background:white; border:1px dashed var(--primary-strong); border-radius:22px; padding:1rem; box-shadow:var(--shadow);}
    [data-testid="stFileUploaderDropzone"] {background:linear-gradient(135deg,var(--soft),var(--secondary-soft)); border-radius:16px; min-height:140px;}
    [data-testid="stMetric"] {
        position:relative; overflow:hidden; min-height:118px; display:flex; justify-content:center;
        background:linear-gradient(145deg,#ffffff 0%,#f4faf6 100%);
        border:1px solid var(--line); border-radius:19px; padding:1.1rem;
        box-shadow:var(--shadow);
    }
    [data-testid="stMetric"]:before {content:""; position:absolute; top:0; right:0; width:100%; height:4px; background:linear-gradient(90deg,var(--primary),var(--secondary));}
    [data-testid="stMetricValue"] {color:var(--primary-dark); font-weight:900; font-size:2rem;}
    [data-testid="stMetricLabel"] {
        color: #4f6c61 !important;
        font-size: .92rem !important;
        font-weight: 700 !important;
        opacity: 1 !important;
        justify-content: center !important;
        text-align: center !important;
        margin-bottom: .35rem;
    }
    [data-testid="stMetricLabel"] p,
    [data-testid="stMetricLabel"] div {
        color: #4f6c61 !important;
        opacity: 1 !important;
    }
    [data-testid="stMetricValue"] > div {
        justify-content: center !important;
        text-align: center !important;
    }
    [data-testid="stDataFrame"] {background:white; border:1px solid var(--line); border-radius:17px; overflow:hidden; box-shadow:0 8px 24px rgba(35,92,72,.06);}
    [data-testid="stExpander"] {background:rgba(255,255,255,.92); border:1px solid var(--line); border-radius:17px; box-shadow:0 7px 20px rgba(35,92,72,.05);}
    [data-testid="stAlert"] {border-radius:16px; border-width:1px; box-shadow:0 6px 18px rgba(35,92,72,.05);}
    [data-testid="stVegaLiteChart"] {background:white; border:1px solid var(--line); border-radius:18px; padding:1rem; box-shadow:var(--shadow);}
    [data-baseweb="tab-list"] {gap:.45rem; background:var(--primary-soft); padding:.35rem; border-radius:14px;}
    [data-baseweb="tab"] {border-radius:10px; color:var(--primary-dark); font-weight:750; padding:.55rem .9rem;}
    [aria-selected="true"][data-baseweb="tab"] {background:white; box-shadow:0 4px 12px rgba(35,92,72,.09);}
    [data-testid="stVerticalBlockBorderWrapper"] {border-color:var(--line) !important; border-radius:22px !important; background:rgba(255,255,255,.72); box-shadow:var(--shadow);}
    .stButton > button, .stDownloadButton > button {
        border-radius:13px; min-height:2.9rem; font-weight:800; border:1px solid var(--primary-strong); transition:all .18s ease;
    }
    .stButton > button:hover {border-color:var(--secondary); color:var(--primary-dark); transform:translateY(-1px);}
    .stDownloadButton > button {background:linear-gradient(110deg,var(--primary-strong),var(--primary-dark)); color:white; width:100%; box-shadow:0 8px 20px rgba(35,92,72,.18);}
    .stDownloadButton > button:hover {background:linear-gradient(110deg,var(--primary-dark),#174635); color:white; border-color:var(--secondary);}
    .privacy-note {background:var(--secondary-soft); border:1px solid #ead49c; color:#70561f; border-radius:15px; padding:.82rem 1rem; font-size:.86rem; margin:.8rem 0 1rem;}
    .step {display:flex; gap:.65rem; align-items:flex-start; margin:.7rem 0; color:var(--muted); font-size:.88rem;}
    .step-num {min-width:27px; height:27px; display:grid; place-items:center; border-radius:50%; background:linear-gradient(135deg,var(--primary),#b9ddca); color:#174b39; border:1px solid rgba(35,92,72,.12); font-weight:900;}
    .analytics-kicker {display:inline-flex; color:var(--secondary-dark); background:var(--secondary-soft); border:1px solid #ead49c; border-radius:99px; padding:.32rem .7rem; font-size:.78rem; font-weight:800; margin-bottom:.65rem;}
    .analytics-note {color:var(--muted); font-size:.86rem; margin:-.15rem 0 .9rem;}
    .donut-card {position:relative; overflow:hidden; min-height:245px; background:linear-gradient(150deg,#fff 0%,#f5faf7 100%); border:1px solid var(--line); border-radius:21px; padding:20px 14px; text-align:center; box-shadow:var(--shadow);}
    .donut-card:before {content:""; position:absolute; top:0; right:0; width:100%; height:4px; background:linear-gradient(90deg,var(--primary),var(--secondary));}
    .donut-ring {width:138px; height:138px; margin:3px auto 14px; border-radius:50%; display:grid; place-items:center; filter:drop-shadow(0 7px 12px rgba(35,92,72,.12));}
    .donut-core {width:98px; height:98px; border-radius:50%; background:#fff; display:grid; place-items:center; color:var(--primary-dark); font-size:1.42rem; font-weight:900; box-shadow:inset 0 0 0 1px #edf4ef;}
    .donut-label {font-weight:900; color:var(--ink); font-size:1rem;}
    .donut-note {font-size:.79rem; color:var(--muted); margin-top:5px;}
    .footer {text-align:center; color:var(--muted); font-size:.82rem; padding:2.6rem 0 1rem; border-top:1px solid var(--line); margin-top:2rem;}
    .footer strong {color:var(--primary-dark);}
    @media (max-width: 760px) {
        [data-testid="stMainBlockContainer"] {padding:1rem .75rem;}
        .hero {padding:1.6rem 1.25rem; border-radius:19px;}
        .feature-grid {grid-template-columns:1fr;}
        [data-testid="stMetric"] {min-height:105px;}
        .donut-card {min-height:225px; margin-bottom:.6rem;}
    }
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown(
    """
    <section class="hero">
      <div class="hero-kicker">EDUCATIONAL QUALITY ANALYTICS</div>
      <h1>منصة تحليلات الجودة التعليمية</h1>
      <p>حوّل ملفات الاستبانات إلى مؤشرات جودة قابلة للقراءة، مع فحص البيانات والتحليل الوصفي وقياس الثبات الداخلي.</p>
      <div class="hero-badges">
        <span class="hero-badge">قراءة ذكية للملفات</span>
        <span class="hero-badge">استبعاد المعرّفات</span>
        <span class="hero-badge">تقارير Excel</span>
      </div>
    </section>
    """,
    unsafe_allow_html=True,
)

with st.sidebar:
    st.markdown("### دليل الاستخدام")
    st.markdown(
        """
        <div class="step"><span class="step-num">1</span><span>ارفع ملف الاستبانة أو ملف النتائج.</span></div>
        <div class="step"><span class="step-num">2</span><span>راجع نوع البيانات الذي اكتشفته المنصة.</span></div>
        <div class="step"><span class="step-num">3</span><span>تحقق من الأسئلة والمعرّفات المستبعدة.</span></div>
        <div class="step"><span class="step-num">4</span><span>راجع المؤشرات ونزّل تقرير Excel.</span></div>
        """,
        unsafe_allow_html=True,
    )
    st.divider()
    st.markdown("#### الصيغ المدعومة")
    st.caption("Excel · CSV · TSV · ODS · PDF · Word · TXT · JSON · XML · Parquet · HTML")
    st.info("أفضل نتيجة تحليلية تتحقق عندما تكون الاستجابات في جدول واضح: صف لكل مشارك وعمود لكل سؤال.")


# ============================================================
# ثوابت عامة
# ============================================================
SUPPORTED_EXTENSIONS = [
    "xlsx", "xls",
    "csv", "tsv",
    "ods",
    "pdf",
    "docx",
    "txt",
    "json",
    "xml",
    "parquet",
    "html", "htm",
]

# حدود تشغيل محافظة للنسخة التجريبية العامة. يمكن رفعها لاحقًا بعد اختبارات أداء.
MAX_UPLOAD_BYTES = 25 * 1024 * 1024
MAX_ROWS = 100_000
MAX_COLUMNS = 500
MAX_TOTAL_CELLS = 10_000_000
MAX_HEADER_LENGTH = 1_000
MAX_DISPLAY_HEADER_LENGTH = 120
MAX_TOTAL_HEADER_CHARS = 100_000
MAX_EXCEL_CELL_CHARS = 32_767
MAX_EXCEL_SHEETS = 30
MAX_PDF_PAGES = 200
MAX_EXTRACTED_TABLES = 100
MAX_ZIP_ENTRIES = 2_000
MAX_UNCOMPRESSED_BYTES = 150 * 1024 * 1024
MAX_COMPRESSION_RATIO = 100

LIKERT_TOKENS = [
    "strongly agree", "agree", "neutral", "disagree", "strongly disagree",
    "very satisfied", "satisfied", "dissatisfied", "very dissatisfied",
    "excellent", "very good", "good", "fair", "poor",
    "موافق بشدة", "موافق", "محايد", "غير موافق", "غير موافق بشدة",
    "راض جدا", "راض جدًا", "راض", "محايد", "غير راض", "غير راض جدا", "غير راض جدًا",
    "ممتاز", "جيد جدا", "جيد جدًا", "جيد", "مقبول", "ضعيف",
]

AGGREGATE_TOKENS = [
    "mean", "average", "avg", "median", "std", "stdev", "sd",
    "variance", "percentage", "percent", "%", "count", "frequency",
    "satisfaction", "rate",
    "المتوسط", "المعدل", "الوسيط", "الانحراف", "التباين",
    "النسبة", "نسبة", "التكرار", "العدد", "الرضا",
]

QUESTION_TOKENS = [
    "question", "item", "statement", "survey item",
    "السؤال", "العبارة", "البند", "الفقرة",
]

IDENTIFIER_TOKENS = [
    "participant_id", "student_id", "respondent_id", "record_id",
    "student_number", "participant_number", "respondent_number",
    "employee_number", "record_number", "serial_number", "serial_no",
    "email", "phone", "mobile", "username", "user_id",
    "معرف", "الرقم الجامعي", "رقم الطالب", "رقم المشارك",
    "رقم الموظف", "الرقم الوظيفي", "الرقم التسلسلي",
    "البريد", "الجوال", "الهاتف", "الاسم",
]


# ============================================================
# أدوات مساعدة
# ============================================================
class UserInputError(ValueError):
    """خطأ متوقع وآمن يمكن عرضه للمستخدم دون كشف تفاصيل الخادم."""


def normalize_text(value) -> str:
    text = str(value).strip().lower()
    text = re.sub(r"\s+", " ", text)
    return text


def user_safe_error(context: str):
    """رسالة عامة لا تكشف مسارات الخادم أو تفاصيل المكتبات."""
    st.error(
        f"تعذر {context}. تحقق من سلامة الملف وبنيته ثم أعد المحاولة. "
        "إذا استمرت المشكلة فاستخدم ملفًا أصغر أو حوّله إلى CSV/XLSX منظم."
    )


def require_optional_access_code():
    """
    بوابة بسيطة للنسخة التجريبية عند ضبط APP_PASSWORD في Streamlit Secrets.
    ليست بديلًا عن نظام هوية وصلاحيات للنسخة التجارية.
    """
    try:
        configured_password = st.secrets.get("APP_PASSWORD")
    except Exception:
        configured_password = None

    if not configured_password:
        st.warning(
            "وضع تجريبي عام: لم يتم تفعيل رمز دخول. استخدم بيانات اصطناعية أو منزوعة الهوية فقط.",
            icon="⚠️",
        )
        return

    if st.session_state.get("access_granted"):
        return

    st.subheader("الدخول إلى النسخة التجريبية")
    entered = st.text_input("رمز الدخول", type="password")
    if st.button("دخول", type="primary", use_container_width=True):
        if hmac.compare_digest(str(entered), str(configured_password)):
            st.session_state["access_granted"] = True
            st.rerun()
        else:
            st.error("رمز الدخول غير صحيح.")
    st.stop()


def validate_zip_container(raw_bytes: bytes, extension: str):
    """تقليل مخاطر ZIP bombs والماكرو داخل صيغ Office المضغوطة."""
    try:
        with zipfile.ZipFile(io.BytesIO(raw_bytes)) as archive:
            infos = archive.infolist()
            if len(infos) > MAX_ZIP_ENTRIES:
                raise UserInputError("يحتوي الملف على عدد مفرط من المكونات الداخلية.")

            total_uncompressed = sum(item.file_size for item in infos)
            total_compressed = sum(max(item.compress_size, 1) for item in infos)
            ratio = total_uncompressed / max(total_compressed, 1)

            if total_uncompressed > MAX_UNCOMPRESSED_BYTES:
                raise UserInputError("الحجم بعد فك الضغط يتجاوز الحد الآمن.")
            if ratio > MAX_COMPRESSION_RATIO:
                raise UserInputError("نسبة الضغط غير طبيعية.")

            names = {item.filename.lower() for item in infos}
            if any(name.endswith("vbaproject.bin") for name in names):
                raise UserInputError("الملفات التي تحتوي وحدات ماكرو غير مسموحة.")

            required_paths = {
                "xlsx": "xl/workbook.xml",
                "docx": "word/document.xml",
            }
            required = required_paths.get(extension)
            if required and required not in names:
                raise UserInputError("البنية الداخلية للملف لا تطابق امتداده.")
    except zipfile.BadZipFile as exc:
        raise UserInputError("بنية الملف المضغوط غير صالحة.") from exc


def validate_uploaded_file(uploaded_file) -> tuple[str, bytes]:
    """التحقق من الحجم والامتداد والبصمة الأولية قبل تمرير الملف إلى القارئ."""
    extension = Path(uploaded_file.name).suffix.lower().lstrip(".")
    if extension not in SUPPORTED_EXTENSIONS:
        raise UserInputError("نوع الملف غير مسموح.")

    uploaded_file.seek(0)
    raw = uploaded_file.read()
    uploaded_file.seek(0)

    if not raw:
        raise UserInputError("الملف فارغ.")
    if len(raw) > MAX_UPLOAD_BYTES:
        raise UserInputError("حجم الملف يتجاوز 25 ميجابايت.")

    signatures = {
        "pdf": lambda b: b.startswith(b"%PDF-"),
        "parquet": lambda b: len(b) >= 8 and b[:4] == b"PAR1" and b[-4:] == b"PAR1",
        "xls": lambda b: b.startswith(bytes.fromhex("D0CF11E0A1B11AE1")),
        "xlsx": lambda b: b.startswith(b"PK\x03\x04"),
        "docx": lambda b: b.startswith(b"PK\x03\x04"),
        "ods": lambda b: b.startswith(b"PK\x03\x04"),
    }
    check = signatures.get(extension)
    if check and not check(raw):
        raise UserInputError("محتوى الملف لا يطابق امتداده.")

    if extension in {"xlsx", "docx", "ods"}:
        validate_zip_container(raw, extension)

    return extension, raw


def enforce_dataframe_limits(df: pd.DataFrame) -> pd.DataFrame:
    if len(df) > MAX_ROWS:
        raise UserInputError(f"عدد الصفوف يتجاوز الحد المسموح ({MAX_ROWS:,}).")
    if len(df.columns) > MAX_COLUMNS:
        raise UserInputError(f"عدد الأعمدة يتجاوز الحد المسموح ({MAX_COLUMNS:,}).")
    total_cells = len(df) * len(df.columns)
    if total_cells > MAX_TOTAL_CELLS:
        raise UserInputError(
            f"حجم الجدول يتجاوز الحد التشغيلي الآمن ({MAX_TOTAL_CELLS:,} خلية). "
            "قسّم الملف إلى أجزاء أصغر."
        )
    return df


def sanitize_excel_value(value):
    """منع تحول نصوص المستخدم إلى صيغ نشطة عند فتح تقرير Excel."""
    if isinstance(value, str):
        cleaned = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]", "", value)
        cleaned = cleaned[:MAX_EXCEL_CELL_CHARS]
        stripped = cleaned.lstrip()
        if stripped.startswith(("=", "+", "-", "@")):
            cleaned = ("'" + cleaned)[:MAX_EXCEL_CELL_CHARS]
        return cleaned
    return value


def sanitize_dataframe_for_excel(df: pd.DataFrame) -> pd.DataFrame:
    safe = df.copy()
    for column in safe.select_dtypes(include=["object", "string"]).columns:
        safe[column] = safe[column].map(sanitize_excel_value)
    safe.columns = [sanitize_excel_value(str(column)) for column in safe.columns]
    return safe


def safe_to_excel(df: pd.DataFrame, writer, sheet_name: str):
    sanitize_dataframe_for_excel(df).to_excel(writer, sheet_name=sheet_name, index=False)


def clean_numeric_series(series: pd.Series) -> pd.Series:
    if pd.api.types.is_numeric_dtype(series):
        return pd.to_numeric(series, errors="coerce").replace([np.inf, -np.inf], np.nan)

    digit_map = str.maketrans("٠١٢٣٤٥٦٧٨٩", "0123456789")

    def normalize_number(value):
        if pd.isna(value):
            return np.nan
        text = str(value).strip().translate(digit_map)
        text = (
            text.replace("٪", "")
            .replace("%", "")
            .replace("−", "-")
            .replace("٬", "")
            .replace("،", ",")
            .replace("٫", ".")
            .replace("\u00a0", "")
            .replace(" ", "")
        )

        # عند وجود الفاصلة والنقطة، نعتبر آخرهما العلامة العشرية.
        if "," in text and "." in text:
            if text.rfind(",") > text.rfind("."):
                text = text.replace(".", "").replace(",", ".")
            else:
                text = text.replace(",", "")
        elif "," in text:
            # 1,000 و1,000,000 فواصل آلاف؛ 2,5 فاصلة عشرية.
            if re.fullmatch(r"[+-]?\d{1,3}(,\d{3})+", text):
                text = text.replace(",", "")
            else:
                text = text.replace(",", ".")

        try:
            number = float(text)
        except (TypeError, ValueError):
            return np.nan
        return number if np.isfinite(number) else np.nan

    return series.map(normalize_number)


def exclude_out_of_range_values(df: pd.DataFrame, columns, minimum: float, maximum: float):
    """إرجاع نسخة منظفة وعدد القيم المستبعدة لكل عمود."""
    cleaned = df.copy()
    excluded = {}
    for col in columns:
        invalid_mask = (
            (cleaned[col] < minimum) | (cleaned[col] > maximum)
        ).fillna(False)
        count = int(invalid_mask.sum())
        if count:
            excluded[col] = count
            cleaned.loc[invalid_mask, col] = np.nan
    return cleaned, excluded


def equal_weight_item_mean(descriptive_df: pd.DataFrame) -> float:
    """المتوسط العام بوزن متساوٍ لكل بند، بصرف النظر عن عدد استجاباته."""
    if descriptive_df.empty or "المتوسط" not in descriptive_df:
        return np.nan
    item_means = pd.to_numeric(descriptive_df["المتوسط"], errors="coerce").dropna()
    return float(item_means.mean()) if not item_means.empty else np.nan


def clean_header_text(value) -> str:
    """تنظيف اسم العمود ومنع العناوين المفرطة أو غير الصالحة لـExcel."""
    name = "" if value is None else str(value)
    name = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]", "", name)
    name = re.sub(r"\s+", " ", name).strip()
    name = name or "Unnamed"
    if len(name) > MAX_HEADER_LENGTH:
        raise UserInputError(
            f"يوجد اسم عمود يتجاوز الحد المسموح ({MAX_HEADER_LENGTH:,} حرف). "
            "اختصر نص السؤال في رأس العمود ثم أعد رفع الملف."
        )
    return name


def make_unique_columns(columns):
    seen = {}
    result = []
    for col in columns:
        name = clean_header_text(col)
        if name not in seen:
            seen[name] = 0
            result.append(name)
        else:
            seen[name] += 1
            suffix = f"_{seen[name]}"
            result.append(f"{name[:MAX_HEADER_LENGTH - len(suffix)]}{suffix}")

    total_chars = sum(len(name) for name in result)
    if total_chars > MAX_TOTAL_HEADER_CHARS:
        raise UserInputError(
            f"إجمالي أحرف أسماء الأعمدة يتجاوز الحد المسموح "
            f"({MAX_TOTAL_HEADER_CHARS:,} حرف)."
        )
    return result


def build_header_aliases(columns):
    """أسماء عرض قصيرة وفريدة مع إبقاء الاسم الكامل للتحليل والتوثيق."""
    aliases = {}
    used = set()
    for index, column in enumerate(columns, start=1):
        full_name = str(column)
        if len(full_name) <= MAX_DISPLAY_HEADER_LENGTH:
            alias = full_name
        else:
            prefix = f"Q{index:03d} · "
            available = MAX_DISPLAY_HEADER_LENGTH - len(prefix) - 1
            alias = f"{prefix}{full_name[:available]}…"

        candidate = alias
        counter = 1
        while candidate in used:
            suffix = f" [{counter}]"
            candidate = f"{alias[:MAX_DISPLAY_HEADER_LENGTH - len(suffix)]}{suffix}"
            counter += 1
        aliases[column] = candidate
        used.add(candidate)
    return aliases


def build_header_mapping(columns, aliases):
    return pd.DataFrame([
        {
            "رقم العمود": index,
            "اسم العرض المختصر": aliases[column],
            "الاسم الأصلي الكامل": str(column),
            "عدد الأحرف": len(str(column)),
            "تم الاختصار": "نعم" if aliases[column] != str(column) else "لا",
        }
        for index, column in enumerate(columns, start=1)
    ])


def rename_columns_for_display(df: pd.DataFrame, aliases):
    return df.rename(columns={column: aliases.get(column, column) for column in df.columns})


def tidy_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df = df.dropna(axis=0, how="all").dropna(axis=1, how="all")

    if df.empty:
        return df

    df.columns = make_unique_columns(df.columns)

    # إزالة أعمدة Unnamed الفارغة غالبًا
    drop_cols = []
    for col in df.columns:
        if normalize_text(col).startswith("unnamed"):
            if df[col].isna().mean() >= 0.95:
                drop_cols.append(col)
    if drop_cols:
        df = df.drop(columns=drop_cols)

    return enforce_dataframe_limits(df.reset_index(drop=True))


def try_promote_first_row_to_header(df: pd.DataFrame) -> pd.DataFrame:
    """
    محاولة ذكية لترقية أول صف إلى أسماء أعمدة عندما تكون الأعمدة الحالية
    أرقامًا أو Unnamed.
    """
    if df.empty or len(df) < 2:
        return df

    current_cols = [normalize_text(c) for c in df.columns]
    suspicious_headers = sum(
        c.startswith("unnamed") or c.isdigit() or c in {"0", "1", "2", "3", "4", "5"}
        for c in current_cols
    )

    first_row = df.iloc[0]
    first_non_null = first_row.notna().sum()
    first_textish = sum(
        isinstance(v, str) and len(v.strip()) > 0
        for v in first_row.dropna().tolist()
    )

    if (
        suspicious_headers >= max(1, len(df.columns) // 2)
        and first_non_null >= max(2, len(df.columns) // 2)
        and first_textish >= max(1, first_non_null // 2)
    ):
        new_df = df.iloc[1:].copy()
        new_df.columns = make_unique_columns(first_row.astype(str).tolist())
        return tidy_dataframe(new_df)

    return df


def detect_encoding(raw_bytes: bytes) -> str:
    for enc in ["utf-8-sig", "utf-8", "cp1256", "windows-1252", "latin1"]:
        try:
            raw_bytes.decode(enc)
            return enc
        except Exception:
            pass
    return "utf-8"


def detect_delimiter(text_sample: str, default=",") -> str:
    try:
        dialect = csv.Sniffer().sniff(text_sample, delimiters=[",", ";", "\t", "|"])
        return dialect.delimiter
    except Exception:
        return default


def looks_like_identifier(column_name: str, series: pd.Series) -> bool:
    name = normalize_text(column_name)
    canonical_name = re.sub(r"[\s-]+", "_", name)

    name_suspicious = (
        canonical_name in IDENTIFIER_TOKENS
        or any(token in name for token in IDENTIFIER_TOKENS if re.search(r"[\u0600-\u06ff]", token))
        or bool(re.search(r"(^|[_\s-])id($|[_\s-])", name))
        or name in {"id", "name", "الاسم", "رقم", "code", "serial"}
    )

    numeric = pd.to_numeric(series, errors="coerce").dropna()

    if numeric.empty:
        return name_suspicious

    unique_ratio = numeric.nunique() / max(len(numeric), 1)
    # لا نعتبر العمود معرفًا فقط لأنه فريد إذا كانت قيمه تبدو كمقياس قصير.
    looks_like_small_scale = (
        numeric.nunique() <= 10
        and numeric.min() >= 0
        and numeric.max() <= 10
    )

    # التسلسل 1،2،3،4،5 وحده ليس دليلًا على أن العمود ID؛ فقد يكون سؤال Likert.
    # نعتمد الاسم الدال على المعرّف، أو التفرد العالي خارج نطاق المقاييس القصيرة.
    return name_suspicious or (unique_ratio >= 0.95 and not looks_like_small_scale)


def is_likert_column_name(column_name: str) -> bool:
    name = normalize_text(column_name)
    return any(token in name for token in LIKERT_TOKENS)


def is_aggregate_column_name(column_name: str) -> bool:
    name = normalize_text(column_name)
    return any(token in name for token in AGGREGATE_TOKENS)


def is_question_column_name(column_name: str) -> bool:
    name = normalize_text(column_name)
    return (
        any(token == name or token in name for token in QUESTION_TOKENS)
        or bool(re.fullmatch(r"q[\s_-]*\d+", name))
    )


def convert_numeric_candidates(df: pd.DataFrame, threshold: float = 0.70) -> pd.DataFrame:
    converted = df.copy()

    for column in converted.columns:
        original_non_missing = converted[column].notna().sum()
        if original_non_missing == 0:
            continue

        numeric_series = clean_numeric_series(converted[column])
        numeric_non_missing = numeric_series.notna().sum()
        ratio = numeric_non_missing / original_non_missing

        if ratio >= threshold:
            converted[column] = numeric_series

    return converted


# ============================================================
# قراءة الملفات
# ============================================================
def read_excel_like(uploaded_file, extension: str):
    uploaded_file.seek(0)

    try:
        if extension == "ods":
            excel = pd.ExcelFile(uploaded_file, engine="odf")
        elif extension == "xls":
            excel = pd.ExcelFile(uploaded_file, engine="xlrd")
        else:
            excel = pd.ExcelFile(uploaded_file)
    except ImportError as exc:
        raise UserInputError(
            "مكتبة قراءة هذا النوع غير مثبتة على الخادم. "
            "راجع ملف requirements.txt ثم أعد نشر المنصة."
        ) from exc

    if len(excel.sheet_names) > MAX_EXCEL_SHEETS:
        raise ValueError(f"عدد أوراق العمل يتجاوز الحد المسموح ({MAX_EXCEL_SHEETS}).")

    candidates = {}
    for sheet in excel.sheet_names:
        try:
            df = pd.read_excel(excel, sheet_name=sheet)
            df = tidy_dataframe(df)
            df = try_promote_first_row_to_header(df)
            if not df.empty:
                candidates[f"Sheet: {sheet}"] = df
        except Exception:
            continue

    return candidates, None


def read_csv_like(uploaded_file, extension: str):
    uploaded_file.seek(0)
    raw = uploaded_file.read()
    encoding = detect_encoding(raw)
    text = raw.decode(encoding, errors="replace")

    delimiter = "\t" if extension == "tsv" else detect_delimiter(text[:5000], default=",")
    df = pd.read_csv(io.StringIO(text), sep=delimiter)
    df = tidy_dataframe(df)
    df = try_promote_first_row_to_header(df)

    return {"Data": df}, f"Encoding: {encoding} | Delimiter: {repr(delimiter)}"


def read_json_file(uploaded_file):
    uploaded_file.seek(0)
    raw = uploaded_file.read()
    encoding = detect_encoding(raw)
    obj = json.loads(raw.decode(encoding, errors="replace"))

    candidates = {}

    if isinstance(obj, list):
        try:
            candidates["JSON records"] = tidy_dataframe(pd.json_normalize(obj))
        except Exception:
            candidates["JSON data"] = tidy_dataframe(pd.DataFrame(obj))

    elif isinstance(obj, dict):
        # محاولة تحويل القاموس كاملًا
        try:
            df = pd.json_normalize(obj)
            if not df.empty:
                candidates["JSON object"] = tidy_dataframe(df)
        except Exception:
            pass

        # البحث عن قوائم داخلية تصلح كجداول
        for key, value in obj.items():
            if isinstance(value, list) and value:
                try:
                    df = pd.json_normalize(value)
                    if not df.empty:
                        candidates[f"JSON: {key}"] = tidy_dataframe(df)
                except Exception:
                    pass

    if not candidates:
        raise ValueError("تعذر تحويل JSON إلى جدول قابل للتحليل.")

    return candidates, None


def read_xml_file(uploaded_file):
    uploaded_file.seek(0)
    raw = uploaded_file.read()
    lowered = raw.lower()
    if b"<!doctype" in lowered or b"<!entity" in lowered:
        raise ValueError("ملفات XML التي تحتوي DOCTYPE أو ENTITY غير مسموحة.")
    df = pd.read_xml(io.BytesIO(raw))
    df = tidy_dataframe(df)
    return {"XML data": df}, None


def read_parquet_file(uploaded_file):
    try:
        import pyarrow.parquet as pq
    except ImportError as exc:
        raise UserInputError(
            "مكتبة PyArrow المطلوبة لقراءة Parquet غير مثبتة على الخادم."
        ) from exc

    uploaded_file.seek(0)
    parquet_file = pq.ParquetFile(uploaded_file)
    metadata = parquet_file.metadata
    if metadata.num_rows > MAX_ROWS:
        raise UserInputError(f"عدد الصفوف يتجاوز الحد المسموح ({MAX_ROWS:,}).")
    if metadata.num_columns > MAX_COLUMNS:
        raise UserInputError(f"عدد الأعمدة يتجاوز الحد المسموح ({MAX_COLUMNS:,}).")
    if metadata.num_rows * metadata.num_columns > MAX_TOTAL_CELLS:
        raise UserInputError(
            f"حجم جدول Parquet يتجاوز الحد التشغيلي الآمن ({MAX_TOTAL_CELLS:,} خلية)."
        )

    uploaded_file.seek(0)
    df = pd.read_parquet(uploaded_file, engine="pyarrow")
    df = tidy_dataframe(df)
    return {"Parquet data": df}, None


def read_html_file(uploaded_file):
    uploaded_file.seek(0)
    raw = uploaded_file.read()
    encoding = detect_encoding(raw)
    html = raw.decode(encoding, errors="replace")

    tables = pd.read_html(io.StringIO(html))
    if len(tables) > MAX_EXTRACTED_TABLES:
        raise ValueError("عدد جداول HTML يتجاوز الحد المسموح.")
    candidates = {}

    for i, table in enumerate(tables, start=1):
        table = tidy_dataframe(table)
        table = try_promote_first_row_to_header(table)
        if not table.empty:
            candidates[f"HTML table {i}"] = table

    if not candidates:
        raise ValueError("لم يتم العثور على جداول HTML قابلة للتحليل.")

    return candidates, None


def read_docx_file(uploaded_file):
    from docx import Document

    uploaded_file.seek(0)
    doc = Document(uploaded_file)

    candidates = {}

    # الجداول أولًا
    if len(doc.tables) > MAX_EXTRACTED_TABLES:
        raise ValueError("عدد جداول Word يتجاوز الحد المسموح.")

    for i, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])

        if not rows:
            continue

        width = max(len(r) for r in rows)
        rows = [r + [""] * (width - len(r)) for r in rows]

        if len(rows) >= 2:
            header = make_unique_columns(rows[0])
            df = pd.DataFrame(rows[1:], columns=header)
        else:
            df = pd.DataFrame(rows)

        df = tidy_dataframe(df)
        if not df.empty:
            candidates[f"Word table {i}"] = df

    # النص الحر كمرشح إضافي
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if paragraphs:
        text_df = pd.DataFrame({"text": paragraphs})
        candidates["Word text"] = text_df

    if not candidates:
        raise ValueError("ملف Word لا يحتوي على جداول أو نص قابل للاستخراج.")

    return candidates, None


def read_pdf_file(uploaded_file):
    import pdfplumber

    uploaded_file.seek(0)
    pdf_bytes = uploaded_file.read()

    candidates = {}
    all_text = []
    table_counter = 0

    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        if len(pdf.pages) > MAX_PDF_PAGES:
            raise ValueError(f"عدد صفحات PDF يتجاوز الحد المسموح ({MAX_PDF_PAGES}).")
        for page_no, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                all_text.append(f"[Page {page_no}]\n{text}")

            try:
                tables = page.extract_tables()
            except Exception:
                tables = []

            for table in tables:
                if table_counter >= MAX_EXTRACTED_TABLES:
                    raise ValueError("عدد الجداول المستخرجة من PDF يتجاوز الحد المسموح.")
                if not table or len(table) < 2:
                    continue

                clean_rows = []
                for row in table:
                    if row is None:
                        continue
                    clean_rows.append([
                        "" if cell is None else str(cell).strip()
                        for cell in row
                    ])

                if len(clean_rows) < 2:
                    continue

                width = max(len(r) for r in clean_rows)
                clean_rows = [r + [""] * (width - len(r)) for r in clean_rows]

                header = make_unique_columns(clean_rows[0])
                df = pd.DataFrame(clean_rows[1:], columns=header)
                df = tidy_dataframe(df)

                if not df.empty:
                    table_counter += 1
                    candidates[f"PDF table {table_counter} - page {page_no}"] = df

    if all_text:
        # نعرض النص كمرشح، لكن لا نفرض أنه بيانات جدولية.
        lines = []
        for chunk in all_text:
            for line in chunk.splitlines():
                line = line.strip()
                if line:
                    lines.append(line)
        if lines:
            candidates["PDF extracted text"] = pd.DataFrame({"text": lines})

    if not candidates:
        raise ValueError(
            "لم يتم استخراج نص أو جداول من PDF. "
            "قد يكون الملف ممسوحًا ضوئيًا (Scanned PDF) ويحتاج OCR."
        )

    return candidates, (
        "PDF: تم استخراج الجداول والنص المتاح. "
        "الملفات الممسوحة ضوئيًا بالكامل قد تحتاج OCR خارجيًا."
    )


def read_txt_file(uploaded_file):
    uploaded_file.seek(0)
    raw = uploaded_file.read()
    encoding = detect_encoding(raw)
    text = raw.decode(encoding, errors="replace")

    candidates = {}

    # محاولة اعتباره ملفًا جدوليًا
    delimiter = detect_delimiter(text[:5000], default=None)
    if delimiter:
        try:
            df = pd.read_csv(io.StringIO(text), sep=delimiter)
            df = tidy_dataframe(df)
            if df.shape[1] >= 2:
                candidates["TXT table"] = df
        except Exception:
            pass

    # النص الحر
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if lines:
        candidates["TXT text"] = pd.DataFrame({"text": lines})

    if not candidates:
        raise ValueError("ملف TXT فارغ أو غير قابل للقراءة.")

    return candidates, f"Encoding: {encoding}"


def load_file(uploaded_file):
    extension, _ = validate_uploaded_file(uploaded_file)

    if extension in {"xlsx", "xls", "ods"}:
        return read_excel_like(uploaded_file, extension), extension

    if extension in {"csv", "tsv"}:
        return read_csv_like(uploaded_file, extension), extension

    if extension == "json":
        return read_json_file(uploaded_file), extension

    if extension == "xml":
        return read_xml_file(uploaded_file), extension

    if extension == "parquet":
        return read_parquet_file(uploaded_file), extension

    if extension in {"html", "htm"}:
        return read_html_file(uploaded_file), extension

    if extension == "docx":
        return read_docx_file(uploaded_file), extension

    if extension == "pdf":
        return read_pdf_file(uploaded_file), extension

    if extension == "txt":
        return read_txt_file(uploaded_file), extension

    raise ValueError("نوع الملف غير مدعوم.")


# ============================================================
# اكتشاف بنية الاستبانة
# ============================================================
def detect_dataset_type(df: pd.DataFrame):
    """
    الأنواع:
    - respondent_level: كل صف مشارك، وكل عمود سؤال
    - frequency_distribution: صف لكل سؤال وأعمدة لفئات الإجابة
    - aggregated: متوسطات/نسب/انحرافات جاهزة
    - text: نص حر غير جدولي
    """
    if df.empty:
        return "unknown", 0.0, "البيانات فارغة"

    if df.shape[1] == 1:
        return "text", 0.95, "تم العثور على عمود نصي واحد."

    normalized_columns = [normalize_text(c) for c in df.columns]

    likert_count = sum(
        any(token in col for token in LIKERT_TOKENS)
        for col in normalized_columns
    )

    aggregate_count = sum(
        any(token in col for token in AGGREGATE_TOKENS)
        for col in normalized_columns
    )

    question_name_count = sum(is_question_column_name(col) for col in normalized_columns)

    converted = convert_numeric_candidates(df)
    numeric_cols = converted.select_dtypes(include=np.number).columns.tolist()
    non_numeric_cols = [c for c in converted.columns if c not in numeric_cols]

    candidate_questions = []
    for col in numeric_cols:
        if not looks_like_identifier(col, converted[col]):
            nunique = converted[col].nunique(dropna=True)
            values = converted[col].dropna()
            if (
                2 <= nunique <= 15
                and not values.empty
                and values.min() >= -1
                and values.max() <= 10
            ):
                candidate_questions.append(col)

    # التوزيع التكراري يحتاج فئات Likert متعددة مع عمود يصف السؤال/البند.
    # هذا يمنع اعتبار استجابات الأفراد توزيعًا لمجرد أسماء بعض الأعمدة.
    has_question_descriptor = question_name_count >= 1 or bool(non_numeric_cols)
    if likert_count >= 2 and has_question_descriptor:
        return (
            "frequency_distribution",
            min(0.98, 0.75 + 0.05 * likert_count),
            "تم اكتشاف عدة أعمدة تحمل تسميات فئات Likert.",
        )

    # الاستجابات الفردية ذات أعمدة قصيرة النطاق لها الأولوية على كلمات عامة
    # مثل rate وaverage التي قد تظهر كمتغير عادي داخل صف المشارك.
    if len(candidate_questions) >= 2 and len(df) >= 5:
        return (
            "respondent_level",
            min(0.98, 0.75 + 0.03 * len(candidate_questions)),
            "تم اكتشاف عدة أعمدة رقمية قصيرة النطاق تبدو كإجابات أفراد.",
        )

    if question_name_count >= 1 and aggregate_count >= 2:
        return (
            "aggregated",
            min(0.98, 0.75 + 0.05 * aggregate_count),
            "يوجد عمود للسؤال مع أعمدة متوسط/نسبة/تكرار أو مؤشرات مجمعة.",
        )

    if aggregate_count >= 2 and (has_question_descriptor or len(df) < 5):
        return (
            "aggregated",
            min(0.95, 0.70 + 0.05 * aggregate_count),
            "تم اكتشاف عدة أعمدة لمؤشرات إحصائية مجمعة.",
        )

    if len(numeric_cols) >= 2 and len(df) >= 5:
        return (
            "respondent_level",
            0.60,
            "البيانات تبدو جدوليّة وبها عدة أعمدة رقمية، لكن الثقة متوسطة.",
        )

    return "unknown", 0.40, "تعذر تحديد بنية الاستبانة تلقائيًا بثقة كافية."


def dataset_type_label(dataset_type: str) -> str:
    labels = {
        "respondent_level": "استجابات فردية Raw / Respondent-level",
        "frequency_distribution": "توزيع تكراري لفئات الإجابة",
        "aggregated": "نتائج مجمعة Aggregated",
        "text": "نص حر",
        "unknown": "غير محدد",
    }
    return labels.get(dataset_type, dataset_type)


def detect_question_columns(df: pd.DataFrame):
    converted = convert_numeric_candidates(df)
    numeric_columns = converted.select_dtypes(include=np.number).columns.tolist()

    candidates = []
    identifiers = []

    for col in numeric_columns:
        if looks_like_identifier(col, converted[col]):
            identifiers.append(col)
            continue

        values = converted[col].dropna()
        if values.empty:
            continue

        unique_count = values.nunique()
        min_val = values.min()
        max_val = values.max()

        # نطاق شائع لمقاييس Likert وغيرها
        likely_scale = (
            unique_count <= 15
            and min_val >= -1
            and max_val <= 10
        )

        # أو أن اسم العمود يبدو كسؤال
        likely_question_name = (
            is_question_column_name(col)
            or bool(re.match(r"^(q|س)\s*[\d_-]+", normalize_text(col)))
        )

        if likely_scale or likely_question_name:
            candidates.append(col)

    return converted, candidates, identifiers


def infer_scale(df: pd.DataFrame, columns):
    if not columns:
        return 1.0, 5.0, 0.0

    values = pd.concat([df[c] for c in columns], ignore_index=True).dropna()
    if values.empty:
        return 1.0, 5.0, 0.0

    actual_min = float(values.min())
    actual_max = float(values.max())

    common_scales = [
        (1.0, 5.0),
        (1.0, 7.0),
        (0.0, 4.0),
        (0.0, 5.0),
        (1.0, 10.0),
        (0.0, 10.0),
    ]

    # أعلى ثقة فقط عندما تظهر نهايتا المقياس فعلًا.
    for lo, hi in common_scales:
        if np.isclose(actual_min, lo) and np.isclose(actual_max, hi):
            return lo, hi, 0.98

    # ظهور الحد الأدنى يسمح باقتراح منخفض/متوسط الثقة. لا نعتمد الحد الأعلى
    # وحده لأن قيم 2–4 مثلًا قد تخص مقياس 1–5 لا مقياس 0–4.
    for lo, hi in common_scales:
        if actual_min >= lo and actual_max <= hi:
            if np.isclose(actual_min, lo):
                return lo, hi, 0.65

    # إذا كانت القيم 2–4 مثلًا، فالحدود الأصلية غير قابلة للاستنتاج من العينة.
    # نعرض 1–5 كافتراض عملي شائع وبثقة منخفضة ليصححه المستخدم عند الحاجة.
    if actual_min >= 1 and actual_max <= 5:
        return 1.0, 5.0, 0.35
    if actual_min >= 0 and actual_max <= 4:
        return 0.0, 4.0, 0.35

    return actual_min, actual_max, 0.25


# ============================================================
# التحليل الإحصائي
# ============================================================
def cronbach_alpha(df: pd.DataFrame) -> float:
    data = df.dropna(axis=0, how="any").astype(float)

    if data.shape[1] < 2:
        raise ValueError("يلزم اختيار سؤالين على الأقل.")
    if data.shape[0] < 2:
        raise ValueError("لا توجد صفوف مكتملة كافية لحساب الثبات.")

    item_variances = data.var(axis=0, ddof=1)
    total_scores = data.sum(axis=1)
    total_variance = total_scores.var(ddof=1)
    n_items = data.shape[1]

    if np.isclose(total_variance, 0):
        raise ValueError("تباين الدرجة الكلية يساوي صفرًا، لذلك لا يمكن حساب الثبات.")

    alpha = (n_items / (n_items - 1)) * (
        1 - item_variances.sum() / total_variance
    )
    return float(alpha)


def classify_alpha(alpha: float) -> str:
    if alpha >= 0.90:
        return "مرتفع جدًا"
    if alpha >= 0.80:
        return "جيد جدًا"
    if alpha >= 0.70:
        return "مقبول"
    if alpha >= 0.60:
        return "حدّي"
    return "ضعيف"


def build_descriptive(df: pd.DataFrame, columns):
    return pd.DataFrame({
        "السؤال": columns,
        "العدد": [df[c].count() for c in columns],
        "المتوسط": [df[c].mean() for c in columns],
        "الانحراف المعياري": [df[c].std(ddof=1) for c in columns],
        "الوسيط": [df[c].median() for c in columns],
        "أقل قيمة": [df[c].min() for c in columns],
        "أعلى قيمة": [df[c].max() for c in columns],
        "القيم المفقودة": [df[c].isna().sum() for c in columns],
    }).round(4)


def build_quality(df: pd.DataFrame, columns, suspicious_columns):
    rows = []
    for col in columns:
        rows.append({
            "العمود": col,
            "عدد القيم": int(df[col].notna().sum()),
            "القيم المفقودة": int(df[col].isna().sum()),
            "نسبة الفقد %": round(df[col].isna().mean() * 100, 2),
            "أقل قيمة": df[col].min(),
            "أعلى قيمة": df[col].max(),
            "عدد القيم الفريدة": int(df[col].nunique(dropna=True)),
            "يبدو كمعرّف": "نعم" if col in suspicious_columns else "لا",
        })
    return pd.DataFrame(rows)


def build_frequency(df: pd.DataFrame, columns):
    frames = []
    for col in columns:
        counts = (
            df[col]
            .value_counts(dropna=False)
            .sort_index()
            .rename_axis("القيمة")
            .reset_index(name="التكرار")
        )
        counts.insert(0, "السؤال", col)
        counts["النسبة %"] = (counts["التكرار"] / len(df) * 100).round(2)
        frames.append(counts)

    return pd.concat(frames, ignore_index=True) if frames else pd.DataFrame()


def prepare_excel_report(
    summary_df,
    descriptive_df=None,
    quality_df=None,
    frequency_df=None,
    selected_df=None,
    source_df=None,
    header_mapping_df=None,
):
    output = io.BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        safe_to_excel(summary_df, writer, "Summary")

        if descriptive_df is not None and not descriptive_df.empty:
            safe_to_excel(descriptive_df, writer, "Descriptive")

        if quality_df is not None and not quality_df.empty:
            safe_to_excel(quality_df, writer, "Data Quality")

        if frequency_df is not None and not frequency_df.empty:
            safe_to_excel(frequency_df, writer, "Frequencies")

        if selected_df is not None and not selected_df.empty:
            safe_to_excel(selected_df, writer, "Selected Data")

        if source_df is not None and not source_df.empty:
            safe_to_excel(source_df.head(50000), writer, "Source Preview")

        if header_mapping_df is not None and not header_mapping_df.empty:
            safe_to_excel(header_mapping_df, writer, "Column Mapping")

    output.seek(0)
    return output


def show_report_preview(sheets: dict, expanded: bool = True):
    """عرض أوراق التقرير داخل المنصة قبل تنزيل ملف Excel."""
    available = [
        (name, frame)
        for name, frame in sheets.items()
        if frame is not None and isinstance(frame, pd.DataFrame) and not frame.empty
    ]

    if not available:
        return

    with st.expander("معاينة التقرير قبل التنزيل", expanded=expanded):
        st.caption(
            "هذه المعاينة تمثل محتوى أوراق ملف Excel. "
            "قد يقتصر عرض البيانات الكبيرة على أول 100 سجل للحفاظ على سرعة المنصة."
        )
        tabs = st.tabs([name for name, _ in available])
        for tab, (_, frame) in zip(tabs, available):
            with tab:
                st.dataframe(frame.head(100), use_container_width=True)
                if len(frame) > 100:
                    st.info(f"تظهر أول 100 سجلات من أصل {len(frame):,} سجلًا.")


def donut_metric(label: str, value: float, color: str = "#136f63", note: str = ""):
    """بطاقة مؤشر دائري خفيفة لا تحتاج مكتبات رسوم إضافية."""
    safe_value = 0.0 if pd.isna(value) else float(np.clip(value, 0, 100))
    st.markdown(
        f"""
        <div class="donut-card">
          <div class="donut-ring"
               style="background:conic-gradient({color} 0 {safe_value:.2f}%,#e4eee8 {safe_value:.2f}% 100%)">
            <div class="donut-core">{safe_value:.1f}%</div>
          </div>
          <div class="donut-label">{label}</div>
          <div class="donut-note">{note}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


# ============================================================
# واجهة رفع الملف
# ============================================================
require_optional_access_code()

st.markdown(
    """
    <div class="feature-grid">
      <div class="feature-card"><b>① اكتشاف تلقائي</b><span>تحديد بنية الاستبانة ونوع البيانات ودرجة الثقة.</span></div>
      <div class="feature-card"><b>② تحليل موثوق</b><span>وصف إحصائي، مؤشر رضا، جودة بيانات وثبات داخلي.</span></div>
      <div class="feature-card"><b>③ تقرير قابل للتنزيل</b><span>ملف Excel منظم يضم الملخص والنتائج التفصيلية.</span></div>
    </div>
    """,
    unsafe_allow_html=True,
)

uploaded_file = st.file_uploader(
    "ابدأ برفع ملف الاستبانة أو نتائجها — الحد الآمن 25 MB",
    type=SUPPORTED_EXTENSIONS,
    help=(
        "مدعوم: Excel بدون ماكرو، CSV, TSV, ODS, PDF, Word DOCX, TXT, "
        "JSON, XML, Parquet, HTML. الحد الأقصى الذي تقبله المنصة 25 MB."
    ),
)

if not uploaded_file:
    st.session_state.pop("analysis_started_for", None)
    st.markdown(
        '<div class="privacy-note">🔒 تنبيه: تجنب رفع الأسماء والأرقام الجامعية وأي بيانات شخصية غير لازمة للتحليل.</div>',
        unsafe_allow_html=True,
    )
    st.info("بانتظار ملف البيانات للبدء في التحليل.", icon="📂")
    st.stop()

uploaded_digest = hashlib.sha256(uploaded_file.getvalue()).hexdigest()
file_key = f"{uploaded_file.name}:{uploaded_digest}"

st.success(
    f"تم استلام الملف **{uploaded_file.name}** وحجمه {uploaded_file.size / 1024:.1f} كيلوبايت. "
    "اضغط الزر لبدء القراءة.",
    icon="📥",
)

if st.button(
    "ابدأ قراءة وتحليل الملف",
    type="primary",
    use_container_width=True,
    icon="▶️",
):
    st.session_state["analysis_started_for"] = file_key

if st.session_state.get("analysis_started_for") != file_key:
    st.info(
        "لم تبدأ المعالجة بعد. يمكنك التأكد من اسم الملف ثم الضغط على زر البدء.",
        icon="ℹ️",
    )
    st.stop()


# ============================================================
# قراءة الملف
# ============================================================
try:
    with st.spinner("جارٍ التحقق من صيغة الملف وقراءة محتواه..."):
        (loaded_result, extension) = load_file(uploaded_file)
        candidates, reader_note = loaded_result
except UserInputError as exc:
    st.error(str(exc))
    st.stop()
except Exception:
    user_safe_error("قراءة الملف")
    st.stop()

if not candidates:
    st.error("لم يتم العثور على بيانات قابلة للتحليل داخل الملف.")
    st.stop()

st.success(f"تمت قراءة الملف بنجاح: {uploaded_file.name}", icon="✅")
st.toast("اكتملت قراءة الملف واستخراج البيانات.", icon="✅")
st.progress(25, text="المرحلة 1 من 4 — تمت قراءة الملف واستخراج البيانات")

if reader_note:
    st.caption(reader_note)

candidate_names = list(candidates.keys())

if len(candidate_names) > 1:
    # نفضّل افتراضيًا أكبر جدول فعلي بدل كائن metadata من صف واحد.
    default_candidate = max(
        candidate_names,
        key=lambda name: (
            len(candidates[name]) * max(len(candidates[name].columns), 1),
            len(candidates[name]),
        ),
    )
    selected_candidate = st.selectbox(
        "تم العثور على أكثر من جدول/قسم. اختر البيانات المطلوب تحليلها",
        candidate_names,
        index=candidate_names.index(default_candidate),
    )
else:
    selected_candidate = candidate_names[0]

raw_df = tidy_dataframe(candidates[selected_candidate])

st.info(
    f"عثرت المنصة على **{len(raw_df):,} استجابة/سجل** و"
    f"**{len(raw_df.columns):,} أعمدة** في الجزء المحدد.",
    icon="🔎",
)

if raw_df.empty:
    st.error("الجزء المحدد لا يحتوي على بيانات.")
    st.stop()

header_aliases = build_header_aliases(raw_df.columns)
header_mapping_df = build_header_mapping(raw_df.columns, header_aliases)
long_header_count = int((header_mapping_df["تم الاختصار"] == "نعم").sum())

if long_header_count:
    st.info(
        f"اكتشفت المنصة **{long_header_count} عنوانًا طويلًا**. "
        "ستُستخدم أسماء مختصرة في العرض، مع حفظ النص الكامل في ورقة Column Mapping بالتقرير.",
        icon="🏷️",
    )
    with st.expander("خريطة العناوين الطويلة", expanded=False):
        st.dataframe(
            header_mapping_df[header_mapping_df["تم الاختصار"] == "نعم"],
            use_container_width=True,
        )


# ============================================================
# Smart File Analyzer
# ============================================================
detected_type, detection_confidence, detection_reason = detect_dataset_type(raw_df)

st.subheader("المحلل الذكي للملف")

c1, c2, c3, c4 = st.columns(4)
c1.metric("نوع الملف", extension.upper())
c2.metric("الصفوف", f"{len(raw_df):,}")
c3.metric("الأعمدة", len(raw_df.columns))
c4.metric("ثقة الاكتشاف", f"{detection_confidence * 100:.0f}%")

st.info(
    f"**نوع البيانات المكتشف:** {dataset_type_label(detected_type)}  \n{detection_reason}",
    icon="🧠",
)

dataset_options = {
    "استجابات فردية Raw / Respondent-level": "respondent_level",
    "توزيع تكراري لفئات الإجابة": "frequency_distribution",
    "نتائج مجمعة Aggregated": "aggregated",
    "نص حر": "text",
    "غير محدد": "unknown",
}

default_label = next(
    (label for label, value in dataset_options.items() if value == detected_type),
    "غير محدد",
)

override_label = st.selectbox(
    "نوع البيانات",
    list(dataset_options.keys()),
    index=list(dataset_options.keys()).index(default_label),
    help="يمكنك تصحيح اكتشاف المنصة إذا كان غير صحيح.",
)

dataset_type = dataset_options[override_label]

with st.expander("معاينة البيانات الأولية", expanded=True):
    st.dataframe(
        rename_columns_for_display(raw_df.head(30), header_aliases),
        use_container_width=True,
    )

st.progress(50, text="المرحلة 2 من 4 — تم اكتشاف بنية البيانات")


# ============================================================
# التعامل مع النص الحر
# ============================================================
if dataset_type == "text":
    st.warning(
        "تم استخراج النص، لكن لا توجد بنية جدولية مؤكدة تسمح بإجراء التحليل الإحصائي مباشرة."
    )

    text_col = raw_df.columns[0]
    extracted_text = "\n".join(raw_df[text_col].astype(str).tolist())

    st.text_area(
        "النص المستخرج",
        value=extracted_text[:30000],
        height=350,
    )

    summary_df = pd.DataFrame([{
        "اسم الملف": uploaded_file.name,
        "نوع الملف": extension,
        "نوع البيانات": dataset_type_label(dataset_type),
        "عدد الأسطر/السجلات": len(raw_df),
        "ملاحظة": "النص يحتاج تحويلًا إلى جدول أو استخراجًا أكثر تخصصًا قبل التحليل الإحصائي.",
    }])

    text_report_df = rename_columns_for_display(raw_df, header_aliases)
    output = prepare_excel_report(
        summary_df,
        source_df=text_report_df,
        header_mapping_df=header_mapping_df,
    )

    show_report_preview({
        "الملخص": summary_df,
        "المحتوى المستخرج": text_report_df,
        "خريطة الأعمدة": header_mapping_df,
    })

    st.download_button(
        "تنزيل النص/البيانات المستخرجة كـ Excel",
        data=output,
        file_name="extracted_content.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    st.stop()


# ============================================================
# النتائج المجمعة أو التوزيعات الجاهزة
# ============================================================
if dataset_type in {"aggregated", "frequency_distribution"}:
    st.subheader("البيانات المستخرجة")

    converted_df = convert_numeric_candidates(raw_df)
    aggregate_display_df = rename_columns_for_display(converted_df, header_aliases)
    st.dataframe(aggregate_display_df, use_container_width=True)

    numeric_columns = converted_df.select_dtypes(include=np.number).columns.tolist()

    if dataset_type == "frequency_distribution":
        likert_columns = [c for c in converted_df.columns if is_likert_column_name(c)]
        if likert_columns:
            st.info(
                "اكتشفت المنصة أعمدة فئات استجابة Likert: "
                + "، ".join(header_aliases.get(col, col) for col in likert_columns)
            )

    st.info(
        "لن يتم حساب Cronbach's Alpha لهذه البيانات تلقائيًا، "
        "لأن الثبات الداخلي يتطلب عادةً استجابات الأفراد لكل بند وليس المتوسطات أو التكرارات المجمعة."
    )

    # عرض ملخص رقمي للحقول الرقمية الموجودة دون الادعاء أنها أسئلة فردية
    aggregate_summary = pd.DataFrame()
    if numeric_columns:
        aggregate_summary = converted_df[numeric_columns].describe().T.reset_index()
        aggregate_summary = aggregate_summary.rename(columns={"index": "المتغير"})
        aggregate_summary["المتغير"] = aggregate_summary["المتغير"].map(
            lambda value: header_aliases.get(value, value)
        )
        st.subheader("ملخص الحقول الرقمية")
        st.dataframe(aggregate_summary, use_container_width=True)

    summary_df = pd.DataFrame([{
        "اسم الملف": uploaded_file.name,
        "نوع الملف": extension,
        "نوع البيانات": dataset_type_label(dataset_type),
        "عدد السجلات": len(converted_df),
        "عدد الأعمدة": len(converted_df.columns),
        "Cronbach Alpha": "غير محسوب - البيانات مجمعة",
    }])

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        safe_to_excel(summary_df, writer, "Summary")
        safe_to_excel(aggregate_display_df, writer, "Extracted Data")
        if not aggregate_summary.empty:
            safe_to_excel(aggregate_summary, writer, "Numeric Summary")
        safe_to_excel(header_mapping_df, writer, "Column Mapping")

    output.seek(0)

    show_report_preview({
        "الملخص": summary_df,
        "البيانات المستخرجة": aggregate_display_df,
        "الملخص الرقمي": aggregate_summary,
        "خريطة الأعمدة": header_mapping_df,
    })

    st.download_button(
        "تنزيل النتائج المستخرجة",
        data=output,
        file_name="educational_quality_aggregated_analysis.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    st.success("اكتملت معالجة البيانات المجمعة.")
    st.stop()


# ============================================================
# الاستجابات الفردية
# ============================================================
converted_df, auto_question_columns, auto_identifier_columns = detect_question_columns(raw_df)

if auto_identifier_columns:
    st.info(
        "استبعدت المنصة مبدئيًا الأعمدة التالية من أسئلة التحليل لأنها تبدو كمعرّفات: "
        + "، ".join(header_aliases.get(col, col) for col in auto_identifier_columns),
        icon="🛡️",
    )

numeric_columns = converted_df.select_dtypes(include=np.number).columns.tolist()

if not numeric_columns:
    st.error(
        "لم يتم العثور على أعمدة رقمية مناسبة. "
        "تأكد من أن استجابات المشاركين موجودة كقيم رقمية."
    )
    st.stop()

st.subheader("تحديد أسئلة الاستبانة")

if auto_question_columns:
    st.success(
        f"اكتشفت المنصة مبدئيًا **{len(auto_question_columns)} أعمدة** تبدو كأسئلة استبانة.",
        icon="✅",
    )
else:
    st.warning(
        "لم تتمكن المنصة من تحديد الأسئلة بثقة. "
        "اختر الأعمدة يدويًا من القائمة."
    )

selected_columns = st.multiselect(
    "اختر أعمدة أسئلة الاستبانة",
    options=numeric_columns,
    default=(
        auto_question_columns
        if auto_question_columns
        else [c for c in numeric_columns if c not in auto_identifier_columns]
    ),
    format_func=lambda column: header_aliases.get(column, column),
    help="لا تُدخل المعرّفات مثل رقم الطالب أو رقم المشارك ضمن الأسئلة.",
)

if not selected_columns:
    st.warning("اختر عمودًا واحدًا على الأقل.")
    st.stop()

suspicious_columns = [
    col for col in selected_columns
    if looks_like_identifier(col, converted_df[col])
]

if suspicious_columns:
    st.warning(
        "الأعمدة التالية تبدو كمعرّفات أو أرقام تسلسلية: "
        + "، ".join(header_aliases.get(col, col) for col in suspicious_columns)
        + ". يفضّل إزالتها من اختيار الأسئلة."
    )

analysis_df = converted_df[selected_columns].copy()

valid_value_count = int(analysis_df.notna().sum().sum())
if valid_value_count == 0:
    st.error(
        "الأعمدة المختارة لا تحتوي قيمًا رقمية صالحة للتحليل. "
        "غيّر الأعمدة أو راجع تنسيق القيم في الملف."
    )
    st.stop()

st.progress(75, text="المرحلة 3 من 4 — تم اعتماد الأعمدة الجاهزة للتحليل")

# ============================================================
# فحص الجودة
# ============================================================
st.subheader("فحص جودة البيانات")

quality_before_df = build_quality(
    analysis_df,
    selected_columns,
    suspicious_columns,
)
quality_before_display_df = quality_before_df.copy()
quality_before_display_df["العمود"] = quality_before_display_df["العمود"].map(
    lambda value: header_aliases.get(value, value)
)
st.dataframe(quality_before_display_df, use_container_width=True)

auto_min, auto_max, scale_confidence = infer_scale(analysis_df, selected_columns)

with st.expander("خيارات مقياس الاستجابة", expanded=True):
    scale_options = {
        "تلقائي وفق البيانات": (float(auto_min), float(auto_max)),
        "Likert من 1 إلى 5": (1.0, 5.0),
        "Likert من 1 إلى 7": (1.0, 7.0),
        "مقياس من 0 إلى 4": (0.0, 4.0),
        "مقياس من 0 إلى 5": (0.0, 5.0),
        "مقياس من 1 إلى 10": (1.0, 10.0),
        "مخصص": None,
    }
    selected_scale = st.selectbox(
        "اختر مقياس الاستجابة الأصلي",
        list(scale_options.keys()),
        help="اختيار المقياس الأصلي أدق من استنتاجه من عينة قد لا تحتوي الحدين الأدنى والأعلى.",
    )

    sc1, sc2, sc3 = st.columns(3)
    preset = scale_options[selected_scale]
    if preset is None:
        expected_min = sc1.number_input("أقل قيمة متوقعة", value=float(auto_min))
        expected_max = sc2.number_input("أعلى قيمة متوقعة", value=float(auto_max))
    else:
        expected_min, expected_max = preset
        sc1.metric("أقل قيمة متوقعة", f"{expected_min:g}")
        sc2.metric("أعلى قيمة متوقعة", f"{expected_max:g}")

    confidence_to_show = scale_confidence if selected_scale.startswith("تلقائي") else 1.0
    sc3.metric("ثقة تحديد المقياس", f"{confidence_to_show * 100:.0f}%")

st.info(
    f"المقياس المستخدم حاليًا من **{expected_min:g}** إلى **{expected_max:g}**. "
    "صححه يدويًا إذا كان مقياس الاستبانة الأصلي مختلفًا.",
    icon="📏",
)

if expected_min >= expected_max:
    st.error("يجب أن تكون أعلى قيمة أكبر من أقل قيمة.")
    st.stop()

out_of_range = {}

for col in selected_columns:
    invalid_count = int(
        (
            (analysis_df[col] < expected_min)
            | (analysis_df[col] > expected_max)
        )
        .fillna(False)
        .sum()
    )

    if invalid_count:
        out_of_range[col] = invalid_count

if out_of_range:
    st.warning(
        "توجد قيم خارج النطاق المتوقع: "
        + "، ".join([
            f"{header_aliases.get(col, col)}: {count}"
            for col, count in out_of_range.items()
        ])
    )
    range_policy = st.radio(
        "كيف تريد التعامل مع القيم خارج النطاق؟",
        ["استبعادها من الحسابات واعتبارها قيمًا مفقودة", "إيقاف التحليل لتصحيح الملف"],
        horizontal=True,
        help="لا تسمح المنصة بإبقاء القيم المخالفة داخل المتوسط أو الرضا أو الثبات.",
    )
    if range_policy.startswith("إيقاف"):
        st.error("تم إيقاف التحليل حتى تُصحّح القيم خارج المقياس في الملف.")
        st.stop()

    analysis_df, excluded_by_column = exclude_out_of_range_values(
        analysis_df,
        selected_columns,
        expected_min,
        expected_max,
    )
    excluded_out_of_range = sum(excluded_by_column.values())

    st.info(
        f"استبعدت المنصة **{excluded_out_of_range:,} قيمة** مخالفة من جميع الحسابات الإحصائية.",
        icon="🧹",
    )
else:
    excluded_out_of_range = 0

if int(analysis_df.notna().sum().sum()) == 0:
    st.error("لا توجد قيم صالحة داخل نطاق المقياس المختار.")
    st.stop()

# التقرير والحسابات اللاحقة يستخدمان النسخة المنظفة فقط.
quality_df = build_quality(analysis_df, selected_columns, suspicious_columns)
quality_df["العمود"] = quality_df["العمود"].map(
    lambda value: header_aliases.get(value, value)
)


# ============================================================
# التحليل الوصفي
# ============================================================
st.subheader("التحليل الوصفي")

descriptive_df = build_descriptive(analysis_df, selected_columns)
descriptive_df["السؤال"] = descriptive_df["السؤال"].map(
    lambda value: header_aliases.get(value, value)
)
st.dataframe(descriptive_df, use_container_width=True)


# ============================================================
# مؤشر الرضا
# ============================================================
st.subheader("مؤشر الرضا")

scale_span = expected_max - expected_min

if scale_span > 0:
    satisfaction_df = descriptive_df[["السؤال", "المتوسط"]].copy()
    satisfaction_df["مؤشر الرضا %"] = (
        (satisfaction_df["المتوسط"] - expected_min)
        / scale_span
        * 100
    ).clip(0, 100).round(2)

    st.dataframe(satisfaction_df, use_container_width=True)

    chart_df = (
        satisfaction_df
        .set_index("السؤال")[["مؤشر الرضا %"]]
        .sort_values("مؤشر الرضا %", ascending=False)
        .head(20)
    )
    st.caption("مقارنة مؤشر الرضا بين البنود — أعلى 20 بندًا")
    st.bar_chart(chart_df, color="#136f63", horizontal=False)
else:
    satisfaction_df = pd.DataFrame()


# ============================================================
# الثبات الداخلي
# ============================================================
st.subheader("الثبات الداخلي")

alpha_value = None
alpha_status = None

if len(selected_columns) < 2:
    st.info("اختر سؤالين على الأقل لحساب Cronbach's Alpha.")
else:
    try:
        complete_rows = int(analysis_df.dropna(axis=0, how="any").shape[0])
        complete_rows_pct = complete_rows / max(len(analysis_df), 1) * 100
        alpha_value = cronbach_alpha(analysis_df)
        alpha_status = classify_alpha(alpha_value)

        a1, a2, a3, a4 = st.columns(4)
        a1.metric("Cronbach's Alpha", f"{alpha_value:.4f}")
        a2.metric("التصنيف", alpha_status)
        a3.metric("الصفوف المستخدمة", f"{complete_rows:,}")
        a4.metric("نسبة العينة المستخدمة", f"{complete_rows_pct:.1f}%")

        if complete_rows_pct < 80:
            st.warning(
                "حُسب Alpha باستخدام الصفوف المكتملة فقط، وتم استبعاد أكثر من 20% من السجلات. "
                "قد تتأثر النتيجة بالقيم المفقودة؛ راجع نمط الفقد قبل اعتمادها."
            )

        if suspicious_columns:
            st.info(
                "قيمة Alpha الحالية تشمل الأعمدة المختارة. "
                "إذا كان بينها معرف، أزله من اختيار الأسئلة."
            )

        if alpha_value < 0 or alpha_value > 1:
            st.warning(
                "قيمة Alpha خارج النطاق المعتاد 0–1. "
                "قد يكون السبب بنودًا معكوسة غير مصححة أو مشكلة في بنية البيانات."
            )

        st.caption(
            "ملاحظة منهجية: البنود العكسية يجب إعادة ترميزها قبل حساب Alpha؛ "
            "المنصة لا تفترض تلقائيًا أي بند عكسي."
        )

    except Exception:
        st.error(
            "تعذر حساب معامل الثبات. تحقق من اختيار بندين على الأقل، "
            "ومن وجود صفوف مكتملة وتباين كافٍ في الإجابات."
        )


# ============================================================
# توزيع الإجابات
# ============================================================
st.subheader("توزيع الإجابات")

frequency_df = build_frequency(analysis_df, selected_columns)
frequency_df["السؤال"] = frequency_df["السؤال"].map(
    lambda value: header_aliases.get(value, value)
)
st.dataframe(frequency_df, use_container_width=True)


# ============================================================
# ملخص عام
# ============================================================
st.subheader("ملخص التحليل")

# وزن متساوٍ لكل بند؛ لا يحصل السؤال ذو الاستجابات الأكثر على وزن أعلى.
overall_mean = equal_weight_item_mean(descriptive_df)

overall_satisfaction = (
    ((overall_mean - expected_min) / (expected_max - expected_min) * 100)
    if expected_max > expected_min
    else np.nan
)
overall_satisfaction = float(np.clip(overall_satisfaction, 0, 100))

m1, m2, m3, m4 = st.columns(4)
m1.metric("عدد السجلات", f"{len(analysis_df):,}")
m2.metric("عدد الأسئلة", len(selected_columns))
m3.metric("المتوسط العام", f"{overall_mean:.3f}")
m4.metric("الرضا العام", f"{overall_satisfaction:.1f}%")

# ============================================================
# لوحة الإحصاءات المرئية
# ============================================================
st.subheader("لوحة الإحصاءات المرئية")
st.markdown(
    '<span class="analytics-kicker">EXECUTIVE ANALYTICS</span>'
    '<div class="analytics-note">قراءة بصرية مركزة لأهم مؤشرات الجودة والاتساق وتوزيع الإجابات.</div>',
    unsafe_allow_html=True,
)

total_cells = analysis_df.shape[0] * analysis_df.shape[1]
available_cells = int(analysis_df.notna().sum().sum())
completeness_pct = (available_cells / total_cells * 100) if total_cells else 0
alpha_pct = float(np.clip(alpha_value * 100, 0, 100)) if alpha_value is not None else 0

with st.container(border=True):
    circle1, circle2, circle3 = st.columns(3, gap="medium")
    with circle1:
        donut_metric("الرضا العام", overall_satisfaction, "#72b493", "محسوب وفق نطاق المقياس")
    with circle2:
        donut_metric("اكتمال البيانات", completeness_pct, "#c9a24b", "نسبة القيم غير المفقودة")
    with circle3:
        donut_metric(
            "الثبات الداخلي",
            alpha_pct,
            "#4f9274",
            alpha_status if alpha_status else "غير متاح",
        )

    st.markdown("<div style='height:.35rem'></div>", unsafe_allow_html=True)
    visual_tab1, visual_tab2 = st.tabs(["مؤشرات الرضا حسب البند", "التوزيع العام للإجابات"])

    with visual_tab1:
        if not satisfaction_df.empty:
            visual_satisfaction = (
                satisfaction_df
                .nlargest(20, "مؤشر الرضا %")
                .sort_values("مؤشر الرضا %", ascending=True)
                .set_index("السؤال")[["مؤشر الرضا %"]]
            )
            chart_height = max(360, min(760, 42 * len(visual_satisfaction)))
            st.caption("مؤشر الرضا لكل بند — تعرض اللوحة 20 بندًا كحد أقصى لتبقى العناوين مقروءة.")
            st.bar_chart(
                visual_satisfaction,
                color="#72b493",
                horizontal=True,
                height=chart_height,
            )
        else:
            st.info("لا تتوفر بيانات كافية لرسم مؤشرات البنود.")

    with visual_tab2:
        combined_answers = (
            pd.concat([analysis_df[c] for c in selected_columns], ignore_index=True)
            .dropna()
            .value_counts()
            .sort_index()
            .rename_axis("قيمة الإجابة")
            .to_frame("التكرار")
        )
        if not combined_answers.empty:
            st.caption("تجميع إجابات البنود المختارة لإظهار شكل التوزيع العام على المقياس.")
            st.bar_chart(combined_answers, color="#c9a24b", height=410)
        else:
            st.info("لا تتوفر إجابات كافية لرسم التوزيع.")

# ============================================================
# الاستجابة التفسيرية للمنصة
# ============================================================
st.subheader("قراءة المنصة للنتائج")

if overall_satisfaction >= 85:
    satisfaction_label = "مرتفع جدًا"
elif overall_satisfaction >= 70:
    satisfaction_label = "جيد"
elif overall_satisfaction >= 50:
    satisfaction_label = "متوسط"
else:
    satisfaction_label = "منخفض"

response_messages = [
    f"بلغ مؤشر الرضا العام **{overall_satisfaction:.1f}%**، ويقع ضمن مستوى **{satisfaction_label}** وفق التصنيف الوصفي للمنصة.",
    f"بلغ اكتمال البيانات **{completeness_pct:.1f}%**؛ والقيم المفقودة تمثل **{100 - completeness_pct:.1f}%** من الخلايا المختارة.",
]

if alpha_value is not None:
    response_messages.append(
        f"بلغ معامل كرونباخ ألفا **{alpha_value:.3f}** وتصنيفه **{alpha_status}**. "
        "هذا مؤشر للاتساق الداخلي، وليس دليلًا مستقلًا على صدق الأداة."
    )
else:
    response_messages.append("لم يتوفر معامل ثبات داخلي قابل للحساب من الاختيارات الحالية.")

if out_of_range:
    response_messages.append(
        f"رصدت المنصة قيمًا خارج نطاق المقياس في **{len(out_of_range)} أعمدة**؛ "
        "ينبغي مراجعتها قبل اعتماد التقرير."
    )

with st.container(border=True):
    for message in response_messages:
        st.info(message, icon="📌")

if not satisfaction_df.empty:
    best_item = satisfaction_df.loc[satisfaction_df["مؤشر الرضا %"].idxmax()]
    weakest_item = satisfaction_df.loc[satisfaction_df["مؤشر الرضا %"].idxmin()]
    insight_left, insight_right = st.columns(2)
    insight_left.success(
        f"أعلى بند رضا: **{best_item['السؤال']}** — {best_item['مؤشر الرضا %']:.1f}%",
        icon="↗️",
    )
    insight_right.warning(
        f"البند الأكثر حاجة للتحسين: **{weakest_item['السؤال']}** — {weakest_item['مؤشر الرضا %']:.1f}%",
        icon="🎯",
    )


# ============================================================
# تنزيل النتائج
# ============================================================
st.subheader("تنزيل النتائج")

summary_df = pd.DataFrame([{
    "اسم الملف": uploaded_file.name,
    "نوع الملف": extension,
    "القسم/الجدول المختار": selected_candidate,
    "نوع البيانات": dataset_type_label(dataset_type),
    "ثقة اكتشاف النوع": detection_confidence,
    "عدد السجلات": len(raw_df),
    "عدد الأسئلة المختارة": len(selected_columns),
    "الأعمدة المشبوهة": ", ".join(
        header_aliases.get(col, col) for col in suspicious_columns
    ),
    "القيم المستبعدة خارج النطاق": excluded_out_of_range,
    "Cronbach Alpha": alpha_value,
    "تصنيف Alpha": alpha_status,
    "النطاق المتوقع": f"{expected_min} - {expected_max}",
    "المتوسط العام": overall_mean,
    "مؤشر الرضا العام %": overall_satisfaction,
    "منهج المتوسط العام": "متوسط متساوي الأوزان لمتوسطات البنود",
}])

output = io.BytesIO()
analysis_report_df = rename_columns_for_display(analysis_df, header_aliases)

with pd.ExcelWriter(output, engine="openpyxl") as writer:
    safe_to_excel(summary_df, writer, "Summary")
    safe_to_excel(descriptive_df, writer, "Descriptive")
    safe_to_excel(quality_df, writer, "Data Quality")
    safe_to_excel(frequency_df, writer, "Frequencies")
    safe_to_excel(analysis_report_df, writer, "Selected Data")

    if not satisfaction_df.empty:
        safe_to_excel(satisfaction_df, writer, "Satisfaction")
    safe_to_excel(header_mapping_df, writer, "Column Mapping")

output.seek(0)

st.progress(100, text="المرحلة 4 من 4 — اكتمل التحليل وأصبح التقرير جاهزًا")

show_report_preview({
    "الملخص": summary_df,
    "التحليل الوصفي": descriptive_df,
    "جودة البيانات": quality_df,
    "مؤشر الرضا": satisfaction_df,
    "التكرارات": frequency_df,
    "البيانات المختارة": analysis_report_df,
    "خريطة الأعمدة": header_mapping_df,
})

st.download_button(
    "تنزيل تقرير Excel الكامل",
    data=output,
    file_name="educational_quality_analysis.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
)

st.success("اكتمل التحليل وأصبح التقرير جاهزًا للتنزيل.", icon="✅")
st.markdown(
    '<div class="footer">منصة تحليلات الجودة التعليمية · تحليل البيانات لاتخاذ قرارات قابلة للقياس<br><strong>تطوير المهندس أحمد المالكي</strong></div>',
    unsafe_allow_html=True,
)
